"""
services/parser.py
====================
Responsible for STEP 4 and STEP 5 of the pipeline:
  - Reading a resume file (PDF or DOCX) from disk
  - Extracting raw text from it

Why this exists:
-----------------
Keeps file-format-specific parsing logic isolated from extraction logic
(extractor.py) and from the route layer (upload.py). Routes should never
know *how* a PDF is read — they just call `parse_resume(file_path)`.
"""

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from app.utils.exceptions import ResumeParsingError


def parse_resume(file_path: Path) -> str:
    """
    Reads a resume file and returns its raw extracted text.

    Args:
        file_path: Absolute path to the saved resume file on disk.

    Returns:
        The full raw text content of the resume.

    Raises:
        ResumeParsingError: if the file type is unsupported, the file
            is corrupted/unreadable, or contains no extractable text
            (e.g. a scanned image-only PDF with no text layer).
    """
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        text = _parse_pdf(file_path)
    elif extension == ".docx":
        text = _parse_docx(file_path)
    else:
        # Defensive check — upload.py should already reject this earlier,
        # but parser.py shouldn't assume it's only ever called safely.
        raise ResumeParsingError(f"Unsupported file extension: {extension}")

    # Guard against PDFs/DOCX that "parsed successfully" but yielded
    # nothing useful (e.g. scanned/image-only resumes with no text layer).
    # We'd rather fail loudly here than silently create an empty candidate.
    if not text or not text.strip():
        raise ResumeParsingError(
            "No extractable text found in the resume. "
            "The file may be a scanned image without a text layer."
        )

    return text.strip()


def _parse_pdf(file_path: Path) -> str:
    """
    Extracts text from a PDF using PyMuPDF (fitz).

    PyMuPDF is used over alternatives (PyPDF2, pdfplumber) because it
    preserves reading order more reliably across multi-column resume
    layouts, which matters a lot for regex-based field extraction later
    (e.g. keeping "Email:" and its value on a coherent line).
    """
    try:
        text_parts: list[str] = []
        # Using a context manager ensures the file handle is released
        # immediately after reading, rather than relying on garbage
        # collection — important since upload.py may later try to
        # delete this same file (on candidate deletion) and a lingering
        # open handle would cause a PermissionError on Windows.
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts)

    except Exception as exc:
        # fitz raises various internal exceptions for corrupted/encrypted
        # PDFs; we normalize all of them into our own domain exception
        # so upload.py only ever has to catch ResumeParsingError.
        raise ResumeParsingError(f"Failed to parse PDF file: {exc}") from exc


def _parse_docx(file_path: Path) -> str:
    """
    Extracts text from a DOCX file using python-docx.

    Reads both regular paragraphs AND table cells, since many resume
    templates use tables for layout (e.g. a two-column skills/contact
    sidebar), and paragraph-only extraction would silently drop that
    content.
    """
    try:
        document = Document(file_path)
        text_parts: list[str] = []

        # Regular paragraphs (most resume body text)
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Table content (common in templated resumes for layout)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)

    except Exception as exc:
        raise ResumeParsingError(f"Failed to parse DOCX file: {exc}") from exc